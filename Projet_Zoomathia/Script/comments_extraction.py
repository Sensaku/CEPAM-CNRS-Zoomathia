import unicodedata

from docx import Document
from lxml import etree
import pandas as pd
from pprint import pprint
import zipfile
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


# Function to extract all the comments of document(Same as accepted answer)
# Returns a dictionary with comment id as key and comment string as value
def get_document_comments(docxFileName):
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    documentXML = docxZip.read('word/document.xml')
    commentsXML = docxZip.read('word/comments.xml')
    etc = etree.XML(commentsXML)
    comments = etc.xpath('//w:comment', namespaces=ooXMLns)
    for c in comments:
        comment = c.xpath('string(.)', namespaces=ooXMLns)
        comment_id = c.xpath('@w:id', namespaces=ooXMLns)[0]
        comments_dict[comment_id] = comment
    return comments_dict


# Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph,comments_dict):
    comments = []
    for run in paragraph.runs:
        comment_reference = run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id = comment_reference[0].xpath('@w:id', namespaces=ooXMLns)[0]
            comment = comments_dict[comment_id]
            comments.append([comment, comment_id])
    return comments


# Function to fetch all comments with their referenced paragraph
# This will return list like this [{'Paragraph text': [comment 1,comment 2]}]
def comments_with_reference_paragraph(docxFileName):

    document = Document(docxFileName)
    comments_dict = get_document_comments(docxFileName)
    comments_with_their_reference_paragraph = []

    for paragraph in document.paragraphs:
        if comments_dict:
            comments = paragraph_comments(paragraph, comments_dict)
            if comments:
                for c in comments:
                    txt = paragraph.text.split("]]")[1].strip()
                    number = paragraph.text.split("]]")[0].replace("[[", "")

                    comments_with_their_reference_paragraph.append(
                        (number, txt, unicodedata.normalize("NFKD", c[0]).strip(), c[1]))

    return comments_with_their_reference_paragraph, paragraphs_extraction


if __name__=="__main__":
    d = "PLINE-8-annotéIPL.docx"  #filepath for the input document
    #comments_with_reference_paragraph(document)
    doc = Document(d)
    docxZip = zipfile.ZipFile(d)
    print(comments_with_reference_paragraph(d))