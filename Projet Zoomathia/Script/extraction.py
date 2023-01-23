import unicodedata
import re
from glob import glob
import rdflib
from os import getcwd
from time import time

# Corese execute
import atexit
import subprocess
from time import sleep
from py4j.java_gateway import JavaGateway
from py4j.java_collections import JavaMap

# Word manipulation
import win32com.client as win32
from win32com.client import constants
from os import path

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from lxml import etree
import zipfile
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


# Start java gateway
java_process = subprocess.Popen(
    ['java', '-jar', '-Dfile.encoding=UTF-8', 'corese-library-python-4.3.0.jar'])
sleep(1)
gateway = JavaGateway()

# Stop java gateway at the enf od script
def exit_handler():
    gateway.shutdown()
    print('\n' * 2)
    print('Gateway Server Stop!')

atexit.register(exit_handler)

# Import of class
Graph = gateway.jvm.fr.inria.corese.core.Graph
Load = gateway.jvm.fr.inria.corese.core.load.Load
Transformer = gateway.jvm.fr.inria.corese.core.transform.Transformer
QueryProcess = gateway.jvm.fr.inria.corese.core.query.QueryProcess
CoreseModel = gateway.jvm.fr.inria.corese.rdf4j.CoreseModel
SimpleValueFactory = gateway.jvm.org.eclipse.rdf4j.model.impl.SimpleValueFactory
RDF = gateway.jvm.org.eclipse.rdf4j.model.vocabulary.RDF

prefLabel = rdflib.SKOS.prefLabel
broader_p = rdflib.SKOS.broader
narrower_p = rdflib.SKOS.narrower
name_file = ""


nb_commentaire, nb_concepts, nb_paragraph = 0, 0, 0
# Function to extract all the comments of document(Same as accepted answer)
# Returns a dictionary with comment id as key and comment string as value
def get_document_comments(docxFileName):
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    documentXML = docxZip.read('word/document.xml')
    commentsXML = docxZip.read('word/comments.xml')
    etc = etree.XML(commentsXML)
    comments = etc.xpath('//w:comment', namespaces=ooXMLns)
    for c in comments:
        comment = c.xpath('string(.)', namespaces=ooXMLns)
        comment_id = c.xpath('@w:id', namespaces=ooXMLns)[0]
        comments_dict[comment_id] = comment
    return comments_dict


# Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph,comments_dict):
    comments = []
    for run in paragraph.runs:
        comment_reference = run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id = comment_reference[0].xpath('@w:id', namespaces=ooXMLns)[0]
            comment = comments_dict[comment_id]
            comments.append([comment, comment_id])
    return comments


# Function to fetch all comments with their referenced paragraph
# This will return list like this [{'Paragraph text': [comment 1,comment 2]}]
def comments_with_reference_paragraph(docxFileName):

    document = Document(docxFileName)
    comments_dict = get_document_comments(docxFileName)
    comments_with_their_reference_paragraph = dict()

    for paragraph in document.paragraphs:
        if comments_dict:
            comments = paragraph_comments(paragraph, comments_dict)
            if comments:
                for c in comments:
                    txt = paragraph.text.split("]]")[1].strip()
                    number = paragraph.text.split("]]")[0].replace("[[", "")

                    comments_with_their_reference_paragraph[c[1]] = [number, txt, unicodedata.normalize("NFKD", c[0]).strip()]

    return comments_with_their_reference_paragraph

def get_comments(filepath):
    global nb_commentaire, nb_concepts
    doc = word.Documents.Open(filepath)
    doc.Activate()
    activeDoc = word.ActiveDocument
    annotation = list()
    nb_commentaire += len(activeDoc.Comments)
    for c in activeDoc.Comments:
        if c.Ancestor is None:  # checking if this is a top-level comment
            for elt in unicodedata.normalize("NFKD", c.Range.Text).strip().split(";"):
                if elt == "":
                    continue
                elt = elt.replace("\n", "").replace("\r", "").replace(u"\xa0", "").strip()
                text = c.Scope.Text.replace("\n", "").replace("\r", "").replace(u"\xa0", "").strip()
                parent = ""
                for candidate in elt.split(":"):
                    data = list()
                    candidate = candidate.replace("\n", "").replace("\r", "").replace(u"\xa0", "").replace("?", "").replace("!","").replace("¿", "").strip()
                    for temp in candidate.split(","):
                        if temp == "":
                            continue
                        nb_concepts += 1
                        for response in skos_exact_search(temp.strip(), parent=parent, commentaire=elt):
                            concept = response.split("idc=")[-1].split("&idt=")[0]
                            if str(c.Index) in extraction_paragraph.keys():
                                annotation.append([c.Index, temp, text, concept, f"{name_file.split('.')[0]}", f"{name_file.split('.')[0].split('-')[1]}", extraction_paragraph[str(c.Index)][0].strip()])

                    parent = candidate
    doc.Close()
    return annotation


def sparqlQuery(graph, query):
    """Run a query on a graph

    :param graph: the graph on which the query is executed
    :param query: query to run
    :returns: query result
    """
    exec = QueryProcess.create(graph)
    return exec.query(query)


def load(path):
    """Load a graph from a local file or a URL

    :param path: local path or a URL
    :returns: the graph load
    """
    graph = Graph()

    ld = Load.create(graph)
    ld.parse(path)

    return graph

def skos_exact_search(label, parent="", commentaire=""):
    if parent != "":
        q = f""" 
    prefix skos: <http://www.w3.org/2004/02/skos/core#>  .

    select distinct * where {{
        ?x skos:prefLabel ?label.
        ?x skos:broader+ ?y .
        ?y skos:narrower+ ?x. 
        ?y skos:prefLabel ?ll.
        filter(contains(str(?ll), "{parent}"))
        filter (str(?label) in (ucase("{label}"), lcase("{label}"), "{label}"))
    }} """
        query_result = sparqlQuery(graph, q)
        candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))
        if candidate == []:
            result = skos_search(label, parent, commentaire)
            if result == []:
                return skos_exact_search(label, parent="", commentaire=commentaire)
            if len(result) > 1:
                many_candidates.append([label, parent, ",".join(result), f"{name_file.split('.')[0]}", commentaire])
            return result
        if len(candidate) > 1:
            many_candidates.append([label, parent, ",".join(candidate), f"{name_file.split('.')[0]}", commentaire])
        return candidate
    else:
        q = f""" 
    prefix skos: <http://www.w3.org/2004/02/skos/core#>  .

    select distinct * where {{
        ?x skos:prefLabel ?label.
        filter (str(?label) in (ucase("{label}"), lcase("{label}"), "{label}"))
    }} """
        query_result = sparqlQuery(graph, q)
        candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))
        if candidate == []:
            return skos_search(label, parent, commentaire)
        if len(candidate) > 1:
            many_candidates.append([label, "", ",".join(candidate), f"{name_file.split('.')[0]}", commentaire])
        return candidate
    return

def skos_search(label, parent="", commentaire=""):

    if parent != "":
        q = f""" 
            prefix skos: <http://www.w3.org/2004/02/skos/core#>  .

            select distinct * where {{
                ?x skos:prefLabel ?label.
                
                ?x skos:broader+ ?y . ?y skos:prefLabel ?ll. 
                ?y skos:narrower+ ?x. ?y skos:prefLabel ?ll.
                filter (contains(str(?ll), "{parent}"))
                filter (contains(str(?label),"{label}"))
            }} """
        query_result = sparqlQuery(graph, q)
        candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))
        if candidate == []:
            return skos_search(label, parent="", commentaire=commentaire)
        if len(candidate) > 1:
            many_candidates.append([label, parent, ",".join(candidate), f"{name_file.split('.')[0]}", commentaire])
        return candidate
    else:
        q = f""" 
            prefix skos: <http://www.w3.org/2004/02/skos/core#>  .

            select distinct * where {{
                ?x skos:prefLabel ?label.
                filter (contains(str(?label), "{label}"))
            }} """
        query_result = sparqlQuery(graph, q)
        candidate = list(set([x.getValue("?x").toString() for x in query_result.getMappingList()]))
        if candidate == []:
            no_concept_found.append([label, "", f"{name_file.split('.')[0]}", commentaire])
        if len(candidate) > 1:
            many_candidates.append([label, "", ", ".join(candidate), f"{name_file.split('.')[0]}", commentaire])
        return candidate
    return


if __name__ == '__main__':
    alpha = time()
    no_concept_found = list()
    many_candidates = list()

    graph = load("th310.ttl")

    files = [x for x in glob("PLINE-*.docx")]

    # Lecture du fichier word: lance un word en arrière plan.

    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False
    paragraph_list = []
    annotation_list = []
    many_candidates = []
    no_concept_found = []

    for file in files:
        print(file)
        name_file = file
        filepath = path.normpath(f"{getcwd()}\\{file}")

        doc_obj = Document(filepath)
        docxZip = zipfile.ZipFile(filepath)

        para_content = [p.text for p in doc_obj.paragraphs]
        nb_paragraph += len(doc_obj.paragraphs)
        paragraph_list.extend([[
            f"{file.split('.')[0]}",
            para_content[0].split("]]")[-1].strip(),
            para_content[1].split("]]")[-1].strip(),
            para_content[2].split("]]")[-1].strip(),
            file.split(".")[0].split("-")[1],
            "".join(p.split("]]")[0].split("[[")[-1]).strip(),
            "".join(p.split("]]")[-1]).strip()] for p in para_content[3:] if p != ""])

        extraction_paragraph = comments_with_reference_paragraph(filepath)
        annotation_list.extend(get_comments(filepath))

    too_many = pd.DataFrame(many_candidates, columns=["label", "parent", "candidate", "chapter", "mention"])
    too_many.to_csv(f"too_many_candidate.csv", index=False, encoding="utf-8", sep=";", line_terminator="\n")

    no_concept = pd.DataFrame(no_concept_found, columns=["label", "parent", "chapter", "mention"])
    no_concept.to_csv(f"no_concept_found.csv", index=False, encoding="utf-8", sep=";", line_terminator="\n")

    extract_para = pd.DataFrame(paragraph_list, columns=["name", "title", "author", "edition", "chapter", "paragraph_number", "paragraphe_text"])
    extract_para.to_csv(f"paragraphs.csv", index=False, encoding="utf-8", sep=";", line_terminator="\n")

    extraction_annotation = pd.DataFrame(annotation_list, columns=["id", "annotation", "text", "concept", "name", "chapter", "paragraph"])
    extraction_annotation.to_csv(f"comment_extraction.csv", index=False, encoding="utf-8", sep=";", line_terminator="\n")

    print(f"Traitement fini en '{time() - alpha}' secondes")
    print(f"nombre de paragraphes {nb_paragraph}, nombre de commentaires {nb_commentaire}, nombre de concepts: {nb_concepts}")
