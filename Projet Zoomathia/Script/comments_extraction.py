from docx import Document
from lxml import etree
import pandas as pd
from pprint import pprint
import zipfile
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


# Function to extract all the comments of document(Same as accepted answer)
# Returns a dictionary with comment id as key and comment string as value
def get_document_comments(docxFileName):
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    with open("test_comments.xml", "w") as tc:
        tc.write(f"{docxZip.read('word/comments.xml')}".replace("'b").replace("\\r\\n"))
    with open("test_document.xml", "w") as td:
        td.write(f"{docxZip.read('word/document.xml')}".replace("'b").replace("\\r\\n"))
    return
    documentXML = docxZip.read('word/document.xml')
    commentsXML = docxZip.read('word/comments.xml')
    etc = etree.XML(commentsXML)
    etd = etree.XML(documentXML)
    comments = etc.xpath('//w:comment',namespaces=ooXMLns)
    #documents = etd.xpath('//w:commentRangeStart/following-sibling::text()[preceding::w:commentRangeEnd]', namespaces=ooXMLns)
    documents = etd.xpath("//w:p/text()[following::@w:commentRangeStart[@w:id=$id] and @w:commentRangeEnd[@w:id=$id]]", namespaces=ooXMLns)
    for c in comments:
        comment = c.xpath('string(.)',namespaces=ooXMLns)
        comment_id = c.xpath('@w:id',namespaces=ooXMLns)[0]
        comments_dict[comment_id] = comment
    return comments_dict


# Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph,comments_dict):
    comments = []
    for run in paragraph.runs:
        comment_reference = run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id = comment_reference[0].xpath('@w:id',namespaces=ooXMLns)[0]
            comment = comments_dict[comment_id]
            comments.append(comment)
    return comments


# Function to fetch all comments with their referenced paragraph
# This will return list like this [{'Paragraph text': [comment 1,comment 2]}]
def comments_with_reference_paragraph(docxFileName):

    document = Document(docxFileName)
    comments_dict = get_document_comments(docxFileName)
    comments_with_their_reference_paragraph = []

    for paragraph in document.paragraphs:
        if comments_dict:
            comments=paragraph_comments(paragraph,comments_dict)
            if comments:
                comments_with_their_reference_paragraph.append({paragraph.text: comments})

    return comments_with_their_reference_paragraph


if __name__=="__main__":
    document = "PLINE-8-annotéIPL.docx"  #filepath for the input document
    #comments_with_reference_paragraph(document)
    doc = Document(document)
    docxZip = zipfile.ZipFile(document)

    documentXML = docxZip.read('word/document.xml')
    commentsXML = docxZip.read('word/comments.xml')
    etc = etree.XML(commentsXML)
    etd = etree.XML(documentXML)
    comments = etc.xpath('//w:comment', namespaces=ooXMLns)
    paragraph = etd.xpath("//w:p", namespaces=ooXMLns)
    for p in paragraph:
        start = p.xpath("w:commentRangeStart", namespaces=ooXMLns)
        end = p.xpath("w:commentRangeEnd", namespaces=ooXMLns)
        for t in text:
            ids = t.xpath("@w:id", namespaces=ooXMLns)
            com = t.xpath("following::w:t", namespaces=ooXMLns)
            print(len(com))
            for c in com:
                tranche = c.xpath(f"w:commentRangeEnd", namespaces=ooXMLns)
                print(tranche)
            print("fin")
            input()